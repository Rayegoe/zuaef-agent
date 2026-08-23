"""Deterministic local report rendering: Markdown -> PDF (reportlab) + DOCX
(python-docx). ADR-008. No Shell, no model call, no business meaning.

The executive-report Skill constrains the supported Markdown subset:
headings, paragraphs, bullet/numbered lists, pipe tables, bold/emphasis,
images, blockquotes, fenced code, horizontal rules. Anything outside this
subset degrades to plain text rather than dropping content.

CJK: PDF text uses the reportlab CID font STSong-Light so Chinese reports
render without external font files.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import mistune
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    Image,
    ListFlowable,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

_PAGE = A4
_MARGIN = 18 * mm
_MAX_IMAGE_WIDTH = _PAGE[0] - 2 * _MARGIN
_MAX_IMAGE_HEIGHT = _PAGE[1] - 2 * _MARGIN

try:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    _BODY_FONT = "STSong-Light"
    _CJK_FONT = True
except Exception:  # noqa: BLE001 — degenerate environment fallback
    _BODY_FONT = "Helvetica"
    _CJK_FONT = False


class RenderError(RuntimeError):
    """Renderer failure with a stable machine code."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


# ── markdown AST walk ───────────────────────────────────────────────────────


def _inline(children: Any) -> str:
    """Render inline AST children to plain text, keeping **strong** and
    *emphasis* markers so the DOCX emitter can build runs."""
    parts: list[str] = []
    for child in children or []:
        kind = child.get("type")
        raw = child.get("raw")
        if kind in ("text", "codespan"):
            parts.append(raw or "")
        elif kind == "emphasis":
            parts.append(f"*{_inline(child.get('children', []))}*")
        elif kind == "strong":
            parts.append(f"**{_inline(child.get('children', []))}**")
        elif kind == "link":
            parts.append(_inline(child.get("children", [])))
        elif kind == "image":
            parts.append(child.get("attrs", {}).get("alt", "") or "")
        elif kind in ("softbreak", "linebreak"):
            parts.append(" ")
        else:
            parts.append(_inline(child.get("children", [])) or raw or "")
    return "".join(parts)


def _blocks_text(blocks: Any) -> str:
    """Inline text of a block container (list items, blockquotes)."""
    parts: list[str] = []
    for block in blocks or []:
        if block.get("type") == "paragraph":
            parts.append(_inline(block.get("children", [])))
        elif block.get("type") == "text":
            parts.append(block.get("raw", ""))
        else:
            parts.append(_blocks_text(block.get("children", [])))
    return " ".join(part for part in parts if part)


def _walk(node: Any) -> list[dict[str, Any]]:
    """Flatten the mistune AST block list into a renderer-neutral walk."""
    out: list[dict[str, Any]] = []
    for block in node:
        kind = block.get("type")
        if kind == "heading":
            out.append(
                {
                    "type": "heading",
                    "level": block.get("attrs", {}).get("level", 1),
                    "children": block.get("children", []),
                }
            )
        elif kind == "paragraph":
            out.append({"type": "paragraph", "children": block.get("children", [])})
        elif kind == "list":
            items: list[str] = []
            for item in block.get("children", []):
                text = _blocks_text(item.get("children", []))
                if text.strip():
                    items.append(text)
            out.append(
                {
                    "type": "list",
                    "ordered": bool(block.get("attrs", {}).get("ordered")),
                    "items": items,
                }
            )
        elif kind == "table":
            rows: list[list[str]] = []
            for row in block.get("children", []):
                cells = row.get("children", [])
                rows.append([_blocks_text([cell]) for cell in cells])
            out.append({"type": "table", "rows": rows})
        elif kind == "blockquote":
            out.append(
                {"type": "blockquote", "text": _blocks_text(block.get("children", []))}
            )
        elif kind == "thematic_break":
            out.append({"type": "rule"})
        elif kind in ("block_code", "code"):
            out.append(
                {
                    "type": "code",
                    "text": (block.get("raw") or block.get("text") or "").rstrip("\n"),
                }
            )
        elif kind == "image":
            attrs = block.get("attrs", {})
            out.append(
                {
                    "type": "image",
                    "src": attrs.get("url", ""),
                    "alt": attrs.get("alt", ""),
                }
            )
        else:
            # Unknown construct: degrade to paragraph text, never drop it.
            out.append(
                {
                    "type": "paragraph",
                    "children": [{"type": "text", "raw": _blocks_text([block])}],
                }
            )
    return out


# ── shared inline rich text ─────────────────────────────────────────────────


def _split_markers(text: str) -> list[tuple[str, bool, bool]]:
    """Split into (segment, bold, italic) runs; unbalanced markers stay text."""
    runs: list[tuple[str, bool, bool]] = []
    bold = italic = False
    buffer = ""
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            if buffer:
                runs.append((buffer, bold, italic))
                buffer = ""
            bold = not bold
            i += 2
        elif text[i] == "*" and not bold:
            if buffer:
                runs.append((buffer, bold, italic))
                buffer = ""
            italic = not italic
            i += 1
        else:
            buffer += text[i]
            i += 1
    if buffer:
        runs.append((buffer, bold, italic))
    return runs


def _pdf_rich(text: str) -> str:
    """Convert **strong**/*emphasis* markers into reportlab markup."""
    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    out: list[str] = []
    for segment, bold, italic in _split_markers(escaped):
        if not segment:
            continue
        wrapped = segment
        if bold:
            wrapped = f"<b>{wrapped}</b>"
        if italic:
            wrapped = f"<i>{wrapped}</i>"
        out.append(wrapped)
    return "".join(out)


# ── PDF ─────────────────────────────────────────────────────────────────────


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "ZBody",
        parent=base["Normal"],
        fontName=_BODY_FONT,
        fontSize=10.5,
        leading=15.5,
        alignment=TA_LEFT,
    )
    return {
        "h1": ParagraphStyle(
            "ZH1",
            parent=base["Heading1"],
            fontName=_BODY_FONT,
            fontSize=17,
            leading=23,
            spaceBefore=14,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "ZH2",
            parent=base["Heading2"],
            fontName=_BODY_FONT,
            fontSize=14,
            leading=19,
            spaceBefore=12,
            spaceAfter=6,
        ),
        "h3": ParagraphStyle(
            "ZH3",
            parent=base["Heading3"],
            fontName=_BODY_FONT,
            fontSize=12,
            leading=16,
            spaceBefore=10,
            spaceAfter=5,
        ),
        "body": body,
        "quote": ParagraphStyle(
            "ZQuote",
            parent=body,
            leftIndent=10 * mm,
            textColor=colors.HexColor("#444444"),
        ),
        "code": ParagraphStyle(
            "ZCode",
            parent=body,
            fontName="Courier",
            fontSize=8.5,
            leading=11,
            leftIndent=6 * mm,
        ),
        "cell": ParagraphStyle(
            "ZCell",
            parent=body,
            fontSize=8.5,
            leading=11.5,
        ),
        "cell_head": ParagraphStyle(
            "ZCellHead",
            parent=body,
            fontSize=8.5,
            leading=11.5,
            textColor=colors.white,
        ),
        "caption": ParagraphStyle(
            "ZCaption",
            parent=body,
            fontSize=8,
            leading=10.5,
            textColor=colors.HexColor("#666666"),
        ),
    }


def _build_pdf_flowables(
    blocks: list[dict[str, Any]], base_dir: Path, styles: dict[str, ParagraphStyle]
) -> list[Any]:
    flowables: list[Any] = []
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            flowables.append(
                Paragraph(
                    _pdf_rich(_inline(block["children"])),
                    styles[f"h{block['level']}"],
                )
            )
        elif kind == "paragraph":
            text = _inline(block["children"])
            if text.strip():
                flowables.append(Paragraph(_pdf_rich(text), styles["body"]))
        elif kind == "list":
            items = [
                Paragraph(_pdf_rich(text), styles["body"]) for text in block["items"]
            ]
            if items:
                flowables.append(
                    ListFlowable(
                        items,
                        bulletType="1" if block["ordered"] else "bullet",
                        start=1 if block["ordered"] else None,
                        leftIndent=14,
                        bulletFontName=_BODY_FONT,
                    )
                )
        elif kind == "table":
            rows = block["rows"]
            if rows:
                data: list[list[Any]] = [
                    [Paragraph(_pdf_rich(cell), styles["cell"]) for cell in row]
                    for row in rows
                ]
                table = Table(data, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                flowables.append(table)
                flowables.append(Spacer(1, 6))
        elif kind == "blockquote":
            flowables.append(Paragraph(_pdf_rich(block["text"]), styles["quote"]))
        elif kind == "rule":
            flowables.append(Spacer(1, 6))
            rule = Table([[""]], colWidths=[_PAGE[0] - 2 * _MARGIN], rowHeights=[1])
            rule.setStyle(
                TableStyle(
                    [("LINEABOVE", (0, 0), (-1, -1), 0.6, colors.HexColor("#aaaaaa"))]
                )
            )
            flowables.append(rule)
            flowables.append(Spacer(1, 6))
        elif kind == "code":
            flowables.append(Spacer(1, 2))
            flowables.append(Preformatted(block["text"], styles["code"]))
            flowables.append(Spacer(1, 4))
        elif kind == "image":
            flowables.append(_pdf_image(block, base_dir, styles))
    return flowables


def _pdf_image(
    block: dict[str, Any], base_dir: Path, styles: dict[str, ParagraphStyle]
) -> Any:
    source = str(block["src"])
    image_file = (base_dir / source).resolve()
    parent = base_dir.resolve()
    if not image_file.is_file() or parent not in image_file.parents:
        return Paragraph(f"[image not found: {source}]", styles["caption"])
    try:
        from PIL import Image as PILImage

        with PILImage.open(image_file) as probe:
            width, height = probe.size
    except Exception:  # noqa: BLE001 — mechanical decode failure
        return Paragraph(f"[image decode failed: {source}]", styles["caption"])
    if width <= 0 or height <= 0:
        return Paragraph(f"[image decode failed: {source}]", styles["caption"])
    scale = min(
        1.0,
        _MAX_IMAGE_WIDTH / width,
        _MAX_IMAGE_HEIGHT / height,
    )
    image = Image(
        str(image_file),
        width=width * scale * 0.35,
        height=height * scale * 0.35,
    )
    if block["alt"]:
        return [image, Paragraph(_pdf_rich(block["alt"]), styles["caption"])]
    return image


def _pdf_footer(canvas: Any, doc: Any) -> None:
    canvas.saveState()
    canvas.setFont(_BODY_FONT, 8)
    canvas.setFillColor(colors.HexColor("#888888"))
    canvas.drawCentredString(_PAGE[0] / 2, 10 * mm, f"{doc.page}")
    canvas.restoreState()


def pdf_page_count(pdf_path: Path) -> int:
    try:
        import pymupdf  # PyMuPDF (also used for preview)

        with pymupdf.open(str(pdf_path)) as document:
            return document.page_count
    except Exception as exc:
        raise RenderError(
            "PDF_UNREADABLE",
            f"rendered PDF could not be opened: {type(exc).__name__}: {exc}",
        ) from exc


def render_pdf(markdown_text: str, pdf_path: Path, base_dir: Path) -> int:
    """Render markdown to a PDF; returns the rendered page count."""
    parser = mistune.create_markdown(renderer=None)
    blocks = _walk(parser(markdown_text))
    styles = _pdf_styles()
    story = _build_pdf_flowables(blocks, base_dir, styles)
    if not story:
        raise RenderError("EMPTY_DOCUMENT", "report.md contains no renderable content")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(pdf_path),
        pagesize=_PAGE,
        leftMargin=_MARGIN,
        rightMargin=_MARGIN,
        topMargin=_MARGIN,
        bottomMargin=_MARGIN,
        title="Competitive Intelligence Report",
        author="ZUAEF",
    )
    document.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    return pdf_page_count(pdf_path)


# ── DOCX ────────────────────────────────────────────────────────────────────


def _docx_runs(paragraph: Any, text: str) -> None:
    for segment, bold, italic in _split_markers(text):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = bold
        run.italic = italic


def render_docx(markdown_text: str, docx_path: Path, base_dir: Path) -> None:
    """Render markdown to an editable DOCX."""
    parser = mistune.create_markdown(renderer=None)
    blocks = _walk(parser(markdown_text))
    document = DocxDocument()
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            _docx_runs(
                document.add_heading("", level=min(block["level"], 3)),
                _inline(block["children"]),
            )
        elif kind == "paragraph":
            text = _inline(block["children"])
            if text.strip():
                _docx_runs(document.add_paragraph(), text)
        elif kind == "list":
            for item in block["items"]:
                if item.strip():
                    _docx_runs(
                        document.add_paragraph(
                            style=("List Number" if block["ordered"] else "List Bullet")
                        ),
                        item,
                    )
        elif kind == "table":
            rows = block["rows"]
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                try:
                    table.style = "Table Grid"
                except Exception:  # noqa: BLE001, S110 — absent style falls back
                    pass
                for row_index, row in enumerate(rows):
                    for col_index, cell in enumerate(row):
                        target = table.cell(row_index, col_index).paragraphs[0]
                        _docx_runs(target, cell)
        elif kind == "blockquote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.4)
            _docx_runs(paragraph, block["text"])
        elif kind == "code":
            for line in block["text"].splitlines():
                paragraph = document.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
        elif kind == "image":
            source = str(block["src"])
            image_file = (base_dir / source).resolve()
            parent = base_dir.resolve()
            if not image_file.is_file() or parent not in image_file.parents:
                continue
            try:
                document.add_picture(str(image_file), width=Inches(5.5))
            except Exception:  # noqa: BLE001, S112 — keep the text report intact
                continue
            if block["alt"]:
                caption = document.add_paragraph()
                caption_run = caption.add_run(block["alt"])
                caption_run.italic = True
                caption_run.font.size = Pt(8)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))


def render_markdown_document(
    markdown_text: str, *, pdf_path: Path, docx_path: Path, base_dir: Path
) -> dict[str, Any]:
    """Full deterministic render; returns mechanical render facts only."""
    pages = render_pdf(markdown_text, pdf_path, base_dir)
    render_docx(markdown_text, docx_path, base_dir)
    return {
        "pages": pages,
        "cjk_font": _CJK_FONT,
        "pdf_path": str(pdf_path),
        "docx_path": str(docx_path),
        "markdown_chars": len(markdown_text),
    }
